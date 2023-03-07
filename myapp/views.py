import os
from io import BytesIO

from django.shortcuts import render, HttpResponse
from django.http import FileResponse

import openai
import docx


def home(request):

    if request.method == 'POST':  # if the request is from the form

        # getting the values submitted in the form
        document_subject = request.POST.get('document_subject')
        document_title = request.POST.get('document_title')
        font = request.POST.get('font')
        length_option = request.POST.get('length_option')
        length_number = request.POST.get('length_number')

        document_length_type = ""
        if length_option == "length_words":
            document_length_type = "words"
        elif length_option == "length_paragraphs":
            document_length_type = "paragraphs"

        gpt_query_length_part = ""
        if length_option != "length_not_specified":
            gpt_query_length_part = ", compose of " + str(length_number) + " " + document_length_type

        gpt_query = "write me " + document_subject + gpt_query_length_part

        openai_api_key = os.environ.get('OPENAI_API_KEY')  # getting my openAI API key from the system variables

        # creating the gpt document
        document = openai.Completion.create(engine="text-davinci-003", prompt=gpt_query, max_tokens=4000)
        document_text = document.choices[0]['text']


        if document_text == "\n\nNone" or document_text == "\n\nSorry, I do not understand what you are asking. Could you please rephrase your question?":
            render(request, 'home.html', {"invalid_query": True})


        # creating a new Word document from the gpt document
        doc = docx.Document()
        if document_title != "":  # if the user set a title, we create it in the new Word document
            doc.add_heading(document_title, 0)
        paragraph = doc.add_paragraph(document_text)  # adding the gpt text to the document
        document_font = paragraph.style.font
        document_font.name = font

        # Save the document to a BytesIO buffer so we can create a FileResponse object out of it
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Create a FileResponse object with the document content
        response = FileResponse(buffer, filename='my_document.docx')
        return response


    else:  # loading the form from the first time
        return render(request, 'home.html')

        